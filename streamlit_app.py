import streamlit as st
import pandas as pd
import altair as alt
from datetime import date
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -----------------------------------------------------------------------------
# GLOBAL STYLING & CONFIG
# -----------------------------------------------------------------------------
st.set_page_config(page_title="Pension Loss Calculator", page_icon="⚖️", layout="wide")

# Hide Streamlit elements to look like a native website tool
st.markdown("""
    <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        .block-container {
            padding-top: 1rem;
            padding-bottom: 0rem;
            padding-left: 1rem;
            padding-right: 1rem;
        }
    </style>
""", unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# OGDEN TABLE DATA MANAGER (DEMO DATASET)
# -----------------------------------------------------------------------------
def get_ogden_subset(gender):
    ages = list(range(40, 61))
    if gender == "Male":
        # Approx 8th Ed multipliers (-0.25%)
        base_60 = [24.50 - (0.95 * (x-40)) for x in ages]
        base_65 = [22.00 - (0.6 * (x-40)) for x in ages]
        base_68 = [19.50 - (0.5 * (x-40)) for x in ages]
        table_name = "Table 28 (Males)"
    else:
        base_60 = [26.00 - (0.90 * (x-40)) for x in ages]
        base_65 = [23.50 - (0.55 * (x-40)) for x in ages]
        base_68 = [21.00 - (0.5 * (x-40)) for x in ages]
        table_name = "Table 29 (Females)"

    df = pd.DataFrame({
        "Age at Trial": ages,
        "Retire at 60": base_60,
        "Retire at 65": base_65,
        "Retire at 68": base_68
    })
    # Ensure no negative numbers
    df = df.applymap(lambda x: max(0.0, x))
    return df, table_name

# -----------------------------------------------------------------------------
# REPORT GENERATOR
# -----------------------------------------------------------------------------
def generate_report(data, results, method):
    doc = Document()
    head = doc.add_heading('Judicial Pension Loss Calculation', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {date.today().strftime('%d %B %Y')}")
    
    p = doc.add_paragraph("Based on: Principles for Compensating Pension Loss (4th Ed, 2021) & Ogden Tables 8th Ed.")
    p.add_runner("\nDISCLAIMER: Draft calculation for estimation only. Uses Term Certain discounting for lump sums.").italic = True

    # 1. Variables
    doc.add_heading('1. Inputs & Assumptions', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    
    def add_row(k, v):
        r = table.add_row().cells
        r[0].text, r[1].text = str(k), str(v)

    if method == "Complex":
        add_row("Claimant Age", data['age'])
        add_row("Target Retirement Age", data['ret_age'])
        add_row("Years to Retirement", f"{data['years_to_retire']} years")
        add_row("Ogden Multiplier (Pension)", f"{data['multiplier']:.2f}")
        add_row("Discount Rate (Lump Sum)", "-0.25% (Term Certain)")
        add_row("Withdrawal (Polkey)", f"{data['withdrawal']}%")
    else:
        add_row("Salary", f"£{data['gross_salary']:,.2f}")

    add_row("Tax Rate", f"{int(data['tax_rate']*100)}%")

    # 2. Calculation
    doc.add_heading('2. Calculation Detail', level=1)
    if method == "Complex":
        doc.add_paragraph(f"Net Annual Loss: £{results['net_annual_loss']:,.2f}")
        doc.add_paragraph(f"Capital Value (Annual): £{results['capital_value_raw']:,.2f}")
        
        # Lump Sum Specifics
        doc.add_heading('Lump Sum Analysis (Accelerated Receipt)', level=2)
        doc.add_paragraph(f"Old Job Lump Sum (Future): £{data['old_lump']:,.2f}")
        doc.add_paragraph(f"New Job Lump Sum (Future): £{data['new_lump_future']:,.2f}")
        doc.add_paragraph(f"Lump Sum Received Early: £{data['new_lump_now']:,.2f}")
        doc.add_paragraph(f"Discount Factor (Term Certain): {results['ls_discount_factor']:.4f}")
        doc.add_paragraph(f"Net Lump Sum Loss (PV): £{results['lump_sum_val']:,.2f}").bold = True
        
        doc.add_heading('Final Totals', level=2)
        doc.add_paragraph(f"Polkey Deduction: -£{results['withdrawal_deduction']:,.2f}")
        doc.add_paragraph(f"Total Net Loss: £{results['net_total']:,.2f}")
    else:
        doc.add_paragraph(f"Total Net Loss: £{results['net_total']:,.2f}")

    # 3. Award
    doc.add_heading('3. Grossed Up Award', level=1)
    doc.add_paragraph(f"Total Award Payable: £{results['gross_total']:,.2f}").bold = True
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------------------------------------------------------
# UI SETUP
# -----------------------------------------------------------------------------
st.title("⚖️ Enhanced Judicial Pension Loss Calculator")
st.markdown("""
**Reference:** *Principles for Compensating Pension Loss (4th Ed, 3rd Rev 2021)*.
Includes **Accelerated Receipt** logic for Lump Sums.
""")

# --- LEGAL DISCLAIMER (Styled to match website) ---
st.markdown("""
    <style>
        .disclaimer-box {
            background-color: #f9f9f9;
            border: 1px solid #e0e0e0;
            padding: 1.2rem;
            color: #666;
            font-size: 0.85rem;
            line-height: 1.4;
            margin-top: 1rem;
            margin-bottom: 2rem;
            border-radius: 0;
        }
    </style>
    <div class="disclaimer-box">
        <strong>For Educational & Illustrative Purposes Only</strong><br><br>
        This is not legal advice. The output is an estimation based on the variables you input. It is not a substitute for a formal opinion from a qualified barrister or solicitor on the specific facts of a case.
        Use of this tool does not create a barrister-client relationship.
        The figures generated are for informational purposes only and should not be relied upon for making any legal, financial, or strategic decisions.
        The accuracy of the output is entirely dependent on the assumptions and data entered. No warranty is given as to its accuracy.
        No liability is accepted for any loss or damage arising from the use of this tool. For professional advice, please seek a formal opinion.
    </div>
""", unsafe_allow_html=True)

# Sidebar
st.sidebar.header("Configuration")
method = st.sidebar.radio("Method", ["Simple (Contributions)", "Complex (Seven Steps)"])
tax_rate = st.sidebar.selectbox("Marginal Tax Rate", [0.20, 0.40, 0.45], index=1, format_func=lambda x: f"{int(x*100)}%")
tax_free_remaining = st.sidebar.number_input("Remaining Tax-Free Allowance (£)", value=0.0)

# -----------------------------------------------------------------------------
# SIMPLE METHOD
# -----------------------------------------------------------------------------
if method == "Simple (Contributions)":
    st.info("Simple method for Defined Contribution schemes.")
    gross_salary = st.number_input("Gross Salary", value=30000.0)
    contrib = st.number_input("Employer %", value=5.0)
    period = st.number_input("Years", value=1.0)
    
    annual_loss = gross_salary * (contrib/100)
    net_total = annual_loss * period
    
    # Gross Up
    taxable = max(0, net_total - tax_free_remaining)
    gross_total = min(net_total, tax_free_remaining) + (taxable / (1 - tax_rate))
    
    st.metric("Total Award", f"£{gross_total:,.2f}")
    
    # Report Data
    r_data = {'gross_salary': gross_salary, 'contrib_rate': contrib, 'tax_rate': tax_rate}
    r_res = {'net_total': net_total, 'gross_total': gross_total}

# -----------------------------------------------------------------------------
# COMPLEX METHOD
# -----------------------------------------------------------------------------
else:
    st.info("Complex method for Defined Benefit / Final Salary schemes.")
    
    # --- STEPS 1-3: ANNUAL PENSION ---
    st.subheader("Steps 1-3: Annual Pension Loss")
    c1, c2, c3 = st.columns(3)
    old_pension = c1.number_input("Projected Old Pension (£)", value=20000.0)
    accrued_pension = c2.number_input("Accrued Old Pension (£)", value=10000.0)
    new_pension = c3.number_input("New Job Pension (£)", value=5000.0)
    
    net_annual_loss = old_pension - (accrued_pension + new_pension)
    st.metric("Net Annual Loss (Multiplicand)", f"£{net_annual_loss:,.2f}")

    # --- STEP 4: MULTIPLIER ---
    st.markdown("---")
    st.subheader("Step 4: Ogden Multiplier")
    
    col_sel1, col_sel2, col_sel3 = st.columns(3)
    gender = col_sel1.selectbox("Gender", ["Male", "Female"])
    claimant_age = col_sel2.number_input("Age at Trial", 40, 60, 50)
    ret_age = col_sel3.selectbox("Retirement Age", [60, 65, 68], index=1)
    
    df_ogden, table_name = get_ogden_subset(gender)
    target_col = f"Retire at {ret_age}"
    
    # Lookup Logic
    try:
        auto_multiplier = df_ogden.loc[df_ogden['Age at Trial'] == claimant_age, target_col].values[0]
        found_in_table = True
    except IndexError:
        auto_multiplier = 0.0
        found_in_table = False

    if found_in_table:
        def highlight_cell(x):
            df1 = pd.DataFrame('', index=x.index, columns=x.columns)
            idx = df_ogden.index[df_ogden['Age at Trial'] == claimant_age].tolist()[0]
            df1.loc[idx, target_col] = 'background-color: #ffeb3b; color: black; font-weight: bold'
            return df1
        st.dataframe(df_ogden.style.apply(highlight_cell, axis=None).format("{:.2f}"), height=200, use_container_width=True)
    else:
        st.warning("Age outside demo range (40-60). Enter manually.")

    c_m1, c_m2 = st.columns(2)
    multiplier = c_m1.number_input("Selected Multiplier", value=float(auto_multiplier), step=0.01)
    withdrawal = c_m2.slider("Polkey Withdrawal (%)", 0, 100, 5)

    # --- LUMP SUMS (WITH ACCELERATED RECEIPT) ---
    st.markdown("---")
    st.subheader("Lump Sums & Accelerated Receipt")
    st.caption("We must separate lump sums received **now** (Accelerated) vs. those received **later** (Future).")

    with st.container():
        col_ls_1, col_ls_2 = st.columns(2)
        
        with col_ls_1:
            st.markdown("##### 1. Old Job (Target)")
            old_lump = st.number_input("Projected Lump Sum at Retirement (£)", value=60000.0, key="old_ls")
            
        with col_ls_2:
            st.markdown("##### 2. Actual Scenario")
            new_lump_future = st.number_input("Future Lump Sum (New Job/Accrued) (£)", value=20000.0, help="Amount you will get at retirement age.")
            new_lump_now = st.number_input("Lump Sum Received Already/Early (£)", value=10000.0, help="E.g., Redundancy pay (excess), early pension release, or refunded contributions.")

    # --- CALCULATION ENGINE ---
    yrs_to_retire = ret_age - claimant_age
    
    # 1. Discount Factor (Term Certain, Ogden Table 35/36 approx)
    # Using Standard Rate -0.25%
    ls_discount_factor = (1 + -0.0025) ** -yrs_to_retire

    # 2. Present Value (PV) Calculations
    # PV of Old Job = Future Amount * Discount
    pv_old_lump = old_lump * ls_discount_factor
    
    # PV of New Job = (Future Amount * Discount) + (Immediate Amount * 1.0)
    pv_new_future = new_lump_future * ls_discount_factor
    pv_new_total = pv_new_future + new_lump_now
    
    # 3. Net Lump Sum Loss
    lump_val = pv_old_lump - pv_new_total

    # Totals
    cap_val = net_annual_loss * multiplier
    total_raw = cap_val + lump_val
    deduction = total_raw * (withdrawal/100)
    net_total = total_raw - deduction
    
    # Gross Up
    taxable = max(0, net_total - tax_free_remaining)
    gross_total = min(net_total, tax_free_remaining) + (taxable / (1 - tax_rate))
    tax_element = gross_total - net_total

    # --- RESULTS ---
    st.markdown("---")
    st.header("Results")
    
    # Detailed Breakdown for Lump Sums
    with st.expander("🔎 View Accelerated Receipt Calculation", expanded=True):
        st.write(f"**Years to Retirement:** {yrs_to_retire} years")
        st.write(f"**Discount Factor (Term Certain):** {ls_discount_factor:.4f}")
        st.markdown("---")
        c_r1, c_r2, c_r3 = st.columns(3)
        c_r1.metric("PV Old Lump Sum", f"£{pv_old_lump:,.2f}", help="Discounted from Future Value")
        c_r2.metric("PV New Lump Sum", f"£{pv_new_total:,.2f}", help=f"Includes £{new_lump_now:,.0f} (undiscounted) + £{pv_new_future:,.0f} (discounted future)")
        c_r3.metric("Net Lump Sum Loss", f"£{lump_val:,.2f}", delta="Step 6 Result")

    st.markdown("#### Final Award")
    m1, m2, m3 = st.columns(3)
    m1.metric("Pension Capital", f"£{cap_val:,.0f}")
    m2.metric("Lump Sum Loss", f"£{lump_val:,.0f}")
    m3.metric("🏆 GROSS AWARD", f"£{gross_total:,.0f}", delta=f"Tax: £{tax_element:,.0f}", delta_color="inverse")

    # --- CHARTS ---
    st.markdown("### 📊 Visual Breakdown")
    chart_col1, chart_col2 = st.columns(2)

    with chart_col1:
        st.caption("Present Value Comparison (Lump Sums)")
        ls_chart_data = pd.DataFrame({
            'Scenario': ['Old Job (PV)', 'Actual (PV)'],
            'Amount': [pv_old_lump, pv_new_total]
        })
        chart1 = alt.Chart(ls_chart_data).mark_bar().encode(
            x='Scenario', y='Amount', color='Scenario'
        ).properties(height=250)
        st.altair_chart(chart1, use_container_width=True)

    with chart_col2:
        st.caption("Final Award Components")
        comp_data = pd.DataFrame({
            'Component': ['Pension Capital', 'Lump Sum Loss', 'Tax Gross-Up'],
            'Value': [cap_val * (1-withdrawal/100), lump_val * (1-withdrawal/100), tax_element]
        })
        chart2 = alt.Chart(comp_data).mark_arc(innerRadius=50).encode(
            theta='Value', color='Component'
        ).properties(height=250)
        st.altair_chart(chart2, use_container_width=True)

    # Report Data Pack
    r_data = {
        'age': claimant_age, 'gender': gender, 'ret_age': ret_age, 
        'table_ref': table_name, 'multiplier': multiplier, 'withdrawal': withdrawal,
        'tax_rate': tax_rate, 'years_to_retire': yrs_to_retire,
        'old_lump': old_lump, 'new_lump_future': new_lump_future, 'new_lump_now': new_lump_now
    }
    r_res = {
        'net_annual_loss': net_annual_loss, 'capital_value_raw': cap_val,
        'lump_sum_val': lump_val, 'withdrawal_deduction': deduction,
        'net_total': net_total, 'gross_total': gross_total,
        'ls_discount_factor': ls_discount_factor
    }

# -----------------------------------------------------------------------------
# DOWNLOAD
# -----------------------------------------------------------------------------
st.markdown("---")
if st.button("Generate Word Report"):
    docx = generate_report(r_data, r_res, "Simple" if method == "Simple (Contributions)" else "Complex")
    st.download_button("Download Report", docx, "pension_report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
